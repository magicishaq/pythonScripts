import docx
#full documentation is from https://python-docx.readthedocs.org
d = docx.Document('C:\\Users\\ikhan\\Documents\\learningPython\\word\\Letter.docx')

for p in d.paragraphs:
    print(p.text)



#each paragraph has a run object, when the text style chages

para = d.paragraphs[0] 
para.runs[0].text

#returns bool
para.runs[0].bold

#set the styleing
p.runs[0].underline = True
#how to save the documnet
# d.save('location')
#p.style

p.style  = 'Title'


 