import docx, os
os.chdir('C:\\Users\\ikhan\\Documents\\learningPython\\word')
string = 'This is a word document created using python :)'
secondString = 'This is another paragraph'
s = docx.Document()
d.add_parapragh(string)

d.add_parapragh(secondString) 


p = d.paragraphs[0] #grabs the first paragraph

p.add_run('This is a new run')
r.runs[0].bold = True #Set the first run  to be bold
d.save('generated.docx')


